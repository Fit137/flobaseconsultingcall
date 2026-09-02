#!/usr/bin/env python3
"""Convert a GTM dashboard markdown file into a formatted DOCX document.

Handles headings, paragraphs, bullet lists, bold spans and markdown tables.

Usage: md_to_docx.py <input.md> <output.docx>
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x38, 0x64)


def is_divider(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def split_row(line):
    return [p.strip() for p in line.strip().strip("|").split("|")]


def add_runs(paragraph, text):
    """Render **bold** and `code` spans inside a paragraph."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def main():
    src, dst = sys.argv[1], sys.argv[2]
    lines = open(src).read().splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and i + 1 < len(lines) and is_divider(lines[i + 1]):
            header = split_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for ci, h in enumerate(header):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                add_runs(cell.paragraphs[0], f"**{re.sub(r'[*]{2}', '', h)}**")
            for row in rows:
                cells = table.add_row().cells
                for ci, v in enumerate(row[: len(header)]):
                    cells[ci].text = ""
                    add_runs(cells[ci].paragraphs[0], v.replace("• ", "\n• "))
            doc.add_paragraph()
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            heading = doc.add_heading(level=min(level, 4))
            run = heading.add_run(text)
            run.font.color.rgb = ACCENT
        elif stripped.startswith("---"):
            doc.add_paragraph()
        elif re.match(r"^[-*] ", stripped):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, stripped)
        i += 1

    doc.save(dst)
    print(f"{dst}: written")


if __name__ == "__main__":
    main()
